from PySimpleGUI import PySimpleGUI as sg
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def doc(cliente, img=0):
    now = datetime.today()
    today = f'{now.day}{now.month}{now.year}--{now.hour}-{now.minute}-{now.second}'
    print(cliente)
    vr = cliente
    document = Document()
    document.add_paragraph(f'Cliente: {vr[0]}')
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(8)

    for v in vr:
        if v == vr[0]: pass
        else:
            check = 0
            print('v = ', v)
            for i in v:
                if check < 1:
                    content = f'''Ficha: {v[0]}
{"Ref.:":<0} {v[1]}      /    Ref. Salto: {v[2]}'''
                    document.add_paragraph(content)
                    if(v[-1] != 0 and v[-1].strip() != ''):
                        my_image = document.add_picture(v[-1], width=Inches(1.5))
                        last_paragraph = document.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                elif check >= 3:
                    if i != v[-1]:
                        content = f'''
{"Cor:":<0} {i[0]}
{"Grade:":<0}  {"33":>10}\t{"34":>10}\t{"35":>10}\t{"36":>10}\t{"37":>10}\t{"38":>10}\t{"39":>10}\t{"40":>10}\t{"Total":>10} 
             {i[1]:>10}\t{i[2]:>10}\t{i[3]:>10}\t{i[4]:>10}\t{i[5]:>10}\t{i[6]:>10}\t{i[7]:>10}\t{i[8]:>10}\t{i[9]:>10}
{"-"*133}'''
                        document.add_paragraph(content)
                check += 1

    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    arq = f'{vr[0]}--{today}.docx'
    caminho = f'{desktop}\{arq}'
    document.save(caminho)


# funções de uso =======================================================================================================

def adicionar(val):
    soma = 0
    if val == 0:
        ficha.append(str(value['ficha']) + '/' + str(value['ficha1']))
        ficha.append(value['ref'])
        if value['refSalto'].split() == ''.split():
            ficha.append('0')
        else:
            ficha.append(value['refSalto'])
        val = 1
    cor.append(value['cor'])
    for i in range(33, 41):
        if value[f'{i}'].split() == ''.split():
            cor.append(0)
        else:
            cor.append(value[f'{i}'])
    for k, v in enumerate(cor):
        if v == '':
            v = 0
        if k > 0:
            soma += int(v)
    cor.append(soma)
    ficha.append(cor.copy())
    cor.clear()
    sg.popup('Adicionado com sucesso')
    return val


# Layout ===========================================================================================================
def janela_cliente():
    sg.theme('Reddit')
    layout = [
        [sg.Text('Cliente'), sg.Input(key='cliente')],
        [sg.Button('Continuar')]
    ]
    return sg.Window('Cliente', layout=layout, finalize=True)

                
def janela_info(nome, reference='', refsalt='', fich='', fich1=''):
    sg.theme('Reddit')
    layout = [
        [sg.Text(f'------------{nome}------------')],
        [sg.Text('Referência:'), sg.Input(default_text=reference, key='ref', size=(10, 1)), sg.Text('Ref. Salto:'), sg.Input(default_text=reference, key='refSalto', size=(10, 1)), sg.Text('Ficha'),
         sg.Input(default_text=fich, key='ficha', size=(5, 1)), sg.Text('/'), sg.Input(default_text=fich1, key='ficha1', size=(5, 1))],
        [sg.Text('Imagem:    '), sg.Button('Foto')],
        [sg.Text('-'*118)],
        [sg.Text('Cor:          '), sg.Input(key='cor', size=(56, 1))],
        [sg.Text('Tamanho:      33      34      35     36     37     38     39     40')],
        [sg.Text('Quatidade: '), sg.Input(key='33', size=(3, 1)), sg.Input(key='34', size=(3, 1)), sg.Input(key='35', size=(3, 1)),
         sg.Input(key='36', size=(3, 1)), sg.Input(key='37', size=(3, 1)), sg.Input(key='38', size=(3, 1)),
         sg.Input(key='39', size=(3, 1)), sg.Input(key='40', size=(3, 1)), sg.Button('Adicionar')],
        [sg.Button('Nova Referência'), sg.Button('Finalizar')]
    ]
    return sg.Window('Informações', layout=layout, finalize=True)


def janela_fim():
    sg.theme('Reddit')
    layout = [
        [sg.Button('Novo Cliente'), sg.Button('Fechar')]
    ]
    return sg.Window('Encerrar', layout=layout, finalize=True)


# Janelas =======================================================================================================
janela1, janela3, janela2, janela4 = janela_cliente(), None, None, None


# Loop de Eventos ===============================================================================================
cliente = []
ficha = []
cor = []
val = 0

while True:
    imagem = ''.strip()
    window, event, value = sg.read_all_windows()
    # fechando
    if event == sg.WINDOW_CLOSED:
        break

    # proxima janela
    if window == janela1 and event == 'Continuar':
        if value['cliente'].strip() == '':
            pass
        else:
            cliente.append(value['cliente'])
            janela2 = janela_info(cliente[0])
            janela1.Close()

    if window == janela2:
        while True:
            window, event, value = sg.read_all_windows()
            if event == sg.WINDOW_CLOSED:
                break

            if event == 'Foto':
                imagem = sg.popup_get_file('Selecione onde está a foto')

            if event == 'Adicionar':
                val = adicionar(val)
                janela2.Close()
                janela2 = janela_info(cliente[0], ficha[1], ficha[2], ficha[0][0], ficha[0][2])

            if event == 'Nova Referência':
                if (imagem == ''.strip() or imagem == None or imagem == 0):
                    try:
                        imagem = sg.popup_get_file('Selecione onde está a foto (deixe vazio para não ter foto)', BaseException=False)
                    except:
                        imagem = ''.strip()
                ficha.append(imagem)
                cliente.append(ficha.copy())
                ficha.clear()
                val = 0
                janela2.Close()
                janela2 = janela_info(cliente[0])
                imagem = ''
                cor.clear()

            if event == 'Finalizar':
                if(imagem == ''.strip() or imagem == None or imagem == 0): 
                    try:
                        imagem = sg.popup_get_file('Selecione onde está a foto (deixe vazio para não ter foto)', BaseException=False)
                    except:
                        imagem = ''.strip()
                ficha.append(imagem)
                cliente.append(ficha.copy())
                ficha.clear()
                doc(cliente, imagem)
                janela4 = janela_fim()
                janela2.Close()
                break


    if window == janela4:
        if event == 'Novo Cliente':
            janela1 = janela_cliente()
            cliente.clear()
            janela4.Close()
            val = 0
        if event == 'Fechar':
            break
